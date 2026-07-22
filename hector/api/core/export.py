"""
Export module — generates PDF and Word documents from HECTOR responses.

Produces court-ready documents with proper formatting, citations,
and metadata. Supports both PDF (via fpdf2) and DOCX (via python-docx).
"""

import io
import logging
import unicodedata
from fpdf.enums import XPos, YPos
from datetime import UTC, datetime

logger = logging.getLogger(__name__)


def _sanitize_for_pdf(text: str) -> str:
    """Replace Unicode characters unsupported by Latin-1 with ASCII equivalents."""
    replacements = {
        "\u2018": "'", "\u2019": "'", "\u201c": '"', "\u201d": '"',
        "\u2014": " - ", "\u2013": "-", "\u2026": "...",
        "\u00a0": " ", "\u2022": "*", "\u00b6": "",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return unicodedata.normalize("NFKD", text).encode("latin-1", "replace").decode("latin-1")


def export_pdf(response_data: dict) -> bytes:
    """
    Generate a PDF document from a HECTOR search response.

    Args:
        response_data: dict with keys: query, generated_response, answer_sections,
                       source_sections, citations, route, confidence, etc.

    Returns:
        PDF file bytes.
    """
    from fpdf import FPDF

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.set_left_margin(10)
    pdf.set_right_margin(10)
    pdf.add_page()
    usable_w = pdf.w - pdf.l_margin - pdf.r_margin

    def _heading(text):
        pdf.set_x(pdf.l_margin)
        pdf.set_font("Helvetica", "B", 11)
        pdf.cell(usable_w, 7, _sanitize_for_pdf(text), new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    def _body(text, size=10):
        pdf.set_x(pdf.l_margin)
        pdf.set_font("Helvetica", "", size)
        pdf.multi_cell(usable_w, 6, _sanitize_for_pdf(text))

    # --- Title ---
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(usable_w, 10, "HECTOR Legal Research Report", new_x=XPos.LMARGIN, new_y=YPos.NEXT, align="C")

    # --- Metadata ---
    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(120, 120, 120)
    generated_at = datetime.now(UTC).strftime("%Y-%m-%d %H:%M UTC")
    pdf.cell(usable_w, 6, f"Generated: {generated_at}", new_x=XPos.LMARGIN, new_y=YPos.NEXT, align="C")

    # --- Divider ---
    pdf.set_draw_color(200, 200, 200)
    pdf.line(10, pdf.get_y(), 200, pdf.get_y())
    pdf.ln(3)

    # --- Query ---
    pdf.set_text_color(0, 0, 0)
    _heading("Query")
    _body(response_data.get("query", "N/A"))

    # --- Route & Confidence ---
    route = response_data.get("route", "LEGAL_RESEARCH")
    confidence = response_data.get("answer_confidence", 0)
    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(100, 100, 100)
    pdf.set_x(pdf.l_margin)
    pdf.cell(usable_w, 5, f"Route: {route}  |  Confidence: {confidence}%", new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    # --- Divider ---
    pdf.set_draw_color(200, 200, 200)
    pdf.line(10, pdf.get_y(), 200, pdf.get_y())
    pdf.ln(3)

    # --- Generated Response ---
    pdf.set_text_color(0, 0, 0)
    _heading("Response")
    response_text = response_data.get("generated_response", "No response generated.")
    for paragraph in response_text.split("\n\n"):
        paragraph = paragraph.strip()
        if paragraph:
                _body(paragraph)

    # --- Answer Sections ---
    answer_sections = response_data.get("answer_sections", [])
    if answer_sections:
        _heading("Answer Sections")
        for i, section in enumerate(answer_sections, 1):
            title = section.get("title", f"Section {i}")
            body = section.get("body", "")
            pdf.set_font("Helvetica", "B", 10)
            pdf.set_x(pdf.l_margin)
            pdf.multi_cell(usable_w, 6, _sanitize_for_pdf(f"{i}. {title}"))
            if body:
                _body(body)

    # --- Citations ---
    citations = response_data.get("citations", [])
    if citations:
        pdf.add_page()
        _heading("Citations")
        pdf.set_font("Helvetica", "", 9)
        for i, citation in enumerate(citations, 1):
            if isinstance(citation, dict):
                act = citation.get("act", "")
                section = citation.get("section", "")
                text = citation.get("text", "")
                cite_str = f"{i}. {act} Section {section}"
                if text:
                    cite_str += f" - {text[:120]}"
            else:
                cite_str = f"{i}. {str(citation)}"
            pdf.set_x(pdf.l_margin)
            pdf.multi_cell(usable_w, 5, _sanitize_for_pdf(cite_str))

    # --- Source Sections ---
    source_sections = response_data.get("source_sections", [])
    if source_sections:
        _heading("Source Sections")
        pdf.set_font("Helvetica", "", 9)
        for i, src in enumerate(source_sections, 1):
            act = src.get("act", "")
            section = src.get("section", "")
            pdf.set_x(pdf.l_margin)
            pdf.multi_cell(usable_w, 5, _sanitize_for_pdf(f"{i}. {act} Section {section}"))

    # --- Footer ---
    pdf.ln(6)
    pdf.set_font("Helvetica", "I", 8)
    pdf.set_text_color(150, 150, 150)
    pdf.set_x(pdf.l_margin)
    pdf.multi_cell(
        usable_w,
        4,
        "This document was generated by HECTOR (Hierarchical Evaluation of "
        "Civil-Criminal Textual's Orchestrator & Retrieval). Responses are "
        "AI-generated and should be verified against official legal texts. "
        "This is not legal advice.",
        align="C",
    )

    # Output bytes
    output = io.BytesIO()
    pdf_data = pdf.output()
    if isinstance(pdf_data, str):
        pdf_data = pdf_data.encode("latin-1")
    output.write(pdf_data)
    return output.getvalue()


def _strip_markdown(text: str) -> str:
    """Remove common markdown formatting from text."""
    import re
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'^#+\s*', '', text, flags=re.MULTILINE)
    text = re.sub(r'^[-*]\s+', '• ', text, flags=re.MULTILINE)
    return text.strip()


def export_docx(response_data: dict) -> bytes:
    """
    Generate a Word document from a HECTOR search response.

    Args:
        response_data: dict with keys: query, generated_response, answer_sections,
                       source_sections, citations, route, confidence, etc.

    Returns:
        DOCX file bytes.
    """
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # --- Style defaults ---
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)

    # --- Title ---
    title = doc.add_heading("HECTOR Legal Research Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- Metadata ---
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"Generated: {datetime.now(UTC).strftime('%Y-%m-%d %H:%M UTC')}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(120, 120, 120)

    # --- Query ---
    doc.add_heading("Query", level=1)
    doc.add_paragraph(response_data.get("query", "N/A"))

    # --- Route & Confidence ---
    route = response_data.get("route", "LEGAL_RESEARCH")
    confidence = response_data.get("answer_confidence", 0)
    meta_p = doc.add_paragraph()
    run = meta_p.add_run(f"Route: {route}  |  Confidence: {confidence}%")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(100, 100, 100)

    # --- Generated Response ---
    doc.add_heading("Response", level=1)
    response_text = response_data.get("generated_response", "No response generated.")
    response_text = _strip_markdown(response_text)
    for paragraph in response_text.split("\n\n"):
        paragraph = paragraph.strip()
        if paragraph:
            doc.add_paragraph(paragraph)

    # --- Answer Sections ---
    answer_sections = response_data.get("answer_sections", [])
    if answer_sections:
        doc.add_heading("Answer Sections", level=1)
        for i, section in enumerate(answer_sections, 1):
            title = section.get("title", f"Section {i}")
            body = _strip_markdown(section.get("body", ""))
            doc.add_heading(f"{i}. {title}", level=2)
            if body:
                doc.add_paragraph(body)
            rows = section.get("rows") or []
            if rows:
                from docx.shared import Inches
                from docx.enum.table import WD_TABLE_ALIGNMENT
                table = doc.add_table(rows=1, cols=3)
                table.style = "Light Grid Accent 1"
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                hdr = table.rows[0].cells
                hdr[0].text = "Point"
                hdr[1].text = "IPC"
                hdr[2].text = "BNS"
                for row in rows:
                    tr = table.add_row().cells
                    tr[0].text = row.get("point", "")
                    tr[1].text = row.get("ipc", "")
                    tr[2].text = row.get("bns", "")
                doc.add_paragraph()

    # --- Citations (filtered: must have act, section, and similarity >= 0.70) ---
    citations = response_data.get("citations", [])
    source_sections = response_data.get("source_sections", [])
    valid_source_acts = {
        s.get("act", "").upper()
        for s in source_sections
        if isinstance(s, dict) and s.get("similarity", 0) >= 0.70
    }
    valid_citations = [
        c for c in citations
        if isinstance(c, dict)
        and c.get("act", "").strip()
        and c.get("section", "").strip()
        and (not valid_source_acts or c.get("act", "").upper() in valid_source_acts)
    ]
    if valid_citations:
        doc.add_heading("Citations", level=1)
        for i, citation in enumerate(valid_citations, 1):
            act = citation.get("act", "")
            section_num = citation.get("section", "")
            text = citation.get("text", "")
            cite_str = f"{act} Section {section_num}"
            if text:
                cite_str += f" — {_strip_markdown(text[:200])}"
            p = doc.add_paragraph()
            run = p.add_run(f"{i}. ")
            run.bold = True
            p.add_run(cite_str)

    # --- Source Sections (filtered: similarity >= 0.70) ---
    valid_sources = [
        s for s in source_sections
        if isinstance(s, dict) and s.get("similarity", 0) >= 0.70
    ]
    if valid_sources:
        doc.add_heading("Source Sections", level=1)
        for i, src in enumerate(valid_sources, 1):
            act = src.get("act", "")
            section_num = src.get("section", "")
            doc.add_paragraph(f"{i}. {act} Section {section_num}")

    # --- Footer ---
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(
        "This document was generated by HECTOR (Hierarchical Evaluation of "
        "Civil-Criminal Textual's Orchestrator & Retrieval). Responses are "
        "AI-generated and should be verified against official legal texts. "
        "This is not legal advice."
    )
    run.font.size = Pt(8)
    run.font.italic = True
    run.font.color.rgb = RGBColor(150, 150, 150)

    # Output bytes
    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()

"""Tests for the PDF and Word export module."""

import os
import sys

import pytest

# Add api/ directory to path so we can import core.export
sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", "..", "api"))

from core.export import export_pdf, export_docx


SAMPLE_RESPONSE = {
    "query": "What is the punishment for murder under BNS?",
    "route": "LEGAL_RESEARCH",
    "generated_response": (
        "Under the Bharatiya Nyaya Sanhita, 2023 (BNS), murder is defined "
        "in Section 101. The punishment for murder is death, or imprisonment "
        "for life, along with a fine.\n\n"
        "This replaces Section 302 of the Indian Penal Code, 1860."
    ),
    "answer_confidence": 85.0,
    "answer_sections": [
        {
            "title": "BNS Section 101 — Murder",
            "body": "Whoever commits murder shall be punished with death, or imprisonment for life, and shall also be liable to fine.",
        }
    ],
    "source_sections": [
        {"act": "BNS", "section": "101"},
        {"act": "IPC", "section": "302"},
    ],
    "citations": [
        {"act": "BNS", "section": "101", "text": "Punishment for murder"},
        {"act": "IPC", "section": "302", "text": "Punishment for murder (repealed)"},
    ],
}


class TestExportPDF:
    def test_returns_bytes(self):
        result = export_pdf(SAMPLE_RESPONSE)
        assert isinstance(result, bytes)

    def test_starts_with_pdf_header(self):
        result = export_pdf(SAMPLE_RESPONSE)
        assert result[:5] == b"%PDF-"

    def test_contains_query(self):
        result = export_pdf(SAMPLE_RESPONSE)
        assert b"punishment for murder" in result.lower()

    def test_contains_disclaimer(self):
        result = export_pdf(SAMPLE_RESPONSE)
        assert b"not legal advice" in result.lower()

    def test_empty_response(self):
        result = export_pdf({})
        assert isinstance(result, bytes)
        assert result[:5] == b"%PDF-"

    def test_with_citations(self):
        result = export_pdf(SAMPLE_RESPONSE)
        assert b"BNS" in result

    def test_with_sections(self):
        result = export_pdf(SAMPLE_RESPONSE)
        assert b"Section 101" in result


class TestExportDocx:
    def test_returns_bytes(self):
        result = export_docx(SAMPLE_RESPONSE)
        assert isinstance(result, bytes)

    def test_is_valid_docx(self):
        result = export_docx(SAMPLE_RESPONSE)
        assert result[:4] == b"PK\x03\x04"

    def test_contains_query(self):
        from docx import Document
        import io

        result = export_docx(SAMPLE_RESPONSE)
        doc = Document(io.BytesIO(result))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "punishment for murder" in text.lower()

    def test_empty_response(self):
        result = export_docx({})
        assert isinstance(result, bytes)
        assert result[:4] == b"PK\x03\x04"

    def test_with_citations(self):
        from docx import Document
        import io

        result = export_docx(SAMPLE_RESPONSE)
        doc = Document(io.BytesIO(result))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "BNS" in text
        assert "IPC" in text

    def test_with_sections(self):
        from docx import Document
        import io

        result = export_docx(SAMPLE_RESPONSE)
        doc = Document(io.BytesIO(result))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Section 101" in text
